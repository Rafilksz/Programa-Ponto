import customtkinter
import mysql.connector
from mysql.connector import Error
import tkinter as tk
from funcoes import center
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from tkinter import filedialog
import datetime
from openpyxl import Workbook
#
#até aqui FUNCIONAVAAAAAAAAAAAAAAAAAA
#até aqui FUNCIONAVAAAAAAAAAAAAAAAAAA
#
#até aqui FUNCIONAVAAAAAAAAAAAAAAAAAA
#até aqui FUNCIONAVAAAAAAAAAAAAAAAAAA
#
#até aqui FUNCIONAVAAAAAAAAAAAAAAAAAA
#
#

customtkinter.set_appearance_mode("dark")
customtkinter.set_default_color_theme("dark-blue")

# Função para converter o nome do mês em um valor numérico
def nome_mes_para_valor(nome_mes):
    meses = {
        "Janeiro": "01",
        "Fevereiro": "02",
        "Março": "03",
        "Abril": "04",
        "Maio": "05",
        "Junho": "06",
        "Julho": "07",
        "Agosto": "08",
        "Setembro": "09",
        "Outubro": "10",
        "Novembro": "11",
        "Dezembro": "12"
    }
    return meses.get(nome_mes, "01")  # Valor padrão é "01" se o nome do mês não for encontrado

def buscar_nomes_funcionarios():
    try:
        con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
        consulta_sql = 'SELECT nome FROM tbl_funcionarios'
        cursor = con.cursor()
        cursor.execute(consulta_sql)
        linhas = cursor.fetchall()
        nomes = [linha[0] for linha in linhas]  # Extrai os nomes da lista de tuplas

        return nomes
    except Error as e:
        print(f"Erro ao buscar nomes dos funcionários: {e}")
    finally:
        if con.is_connected():
            cursor.close()
            con.close()

def buscar_dados_funcionario():
    nome_selecionado = cb_nomes.get()  # Obtém o nome selecionado na combobox
    mes_selecionado = cb_mes.get()  # Obtém o mês selecionado na combobox

    valor_mes = nome_mes_para_valor(mes_selecionado)  # Converte o nome do mês para valor numérico

    try:
        con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
        consulta_sql = f'SELECT * FROM tbl_horario WHERE nomefuncionario = "{nome_selecionado}" AND MONTH(data) = "{valor_mes}"'
        cursor = con.cursor()
        cursor.execute(consulta_sql)
        dados = cursor.fetchall()

        for linha in dados:
            id_funcionario, nomefuncionario, data, ini_exp_time, ini_alm_time, fim_alm_time, fim_exp_time = linha
            linha_formatada = f"Cpf: {id_funcionario},Nome {nomefuncionario}, Data: {data}, Início do Expediente: {ini_exp_time}, Início do Almoço: {ini_alm_time}, Término do Almoço: {fim_alm_time}, Término do Expediente: {fim_exp_time}"
            print(linha_formatada)

    except Error as e:
        print(f"Erro ao buscar dados do funcionário: {e}")
    finally:
        if con.is_connected():
            cursor.close()
            con.close()


def buscar_e_salvar_como_excel():
    nome_selecionado = cb_nomes.get()

    # Obter o mês e o ano atuais
    data_atual = datetime.datetime.now()
    mes_atual = data_atual.month
    ano_atual = data_atual.year

    # Consulta SQL para buscar os dados do funcionário no mês e ano atuais
    consulta_sql = f"SELECT * FROM tbl_horario WHERE nomefuncionario = '{nome_selecionado}' AND YEAR(data) = {ano_atual} AND MONTH(data) = {mes_atual}"

    try:
        con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
        cursor = con.cursor()
        cursor.execute(consulta_sql)
        resultados = cursor.fetchall()

        # Abre uma janela de diálogo para escolher o local de salvamento
        arquivo = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Arquivos Excel", "*.xlsx")])

        if arquivo:
            # Criar um novo arquivo Excel
            workbook = Workbook()
            sheet = workbook.active

            # Definir o tamanho das colunas
            column_sizes = {'A': 15, 'B': 20, 'C': 15, 'D': 20, 'E': 15, 'F': 15, 'G': 15}
            for col, size in column_sizes.items():
                sheet.column_dimensions[col].width = size

            # Adicionar cabeçalhos às colunas, incluindo a coluna "Soma Total"
            headers = ["CPF", "Nome Funcionario", "Data", "Inicio Expediente", "Inicio Almoço", "Fim Almoço",
                       "Fim Expediente", "Soma Total"]
            sheet.append(headers)

            # Adicionar os dados do banco de dados ao arquivo Excel e calcular a soma total
            for resultado in resultados:
                cpf = resultado[0]
                nome_funcionario = resultado[1]
                data = resultado[2]
                inicio_expediente = resultado[3]
                inicio_almoco = resultado[4]
                fim_almoco = resultado[5]
                fim_expediente = resultado[6]  # Supondo que o fim de expediente está na sétima coluna

                # Verificar se os valores não são None antes de calcular o tempo trabalhado
                if inicio_expediente is not None and fim_expediente is not None and inicio_almoco is not None and fim_almoco is not None:
                    # Calcular o tempo trabalhado
                    tempo_trabalhado = (fim_expediente - inicio_expediente) - (fim_almoco - inicio_almoco)
                else:
                    tempo_trabalhado = None

                # Adicionar os dados à linha, incluindo o tempo trabalhado
                linha = [cpf, nome_funcionario, data, inicio_expediente, inicio_almoco, fim_almoco, fim_expediente,
                         tempo_trabalhado]
                sheet.append(linha)

            # Salvar o arquivo Excel no local escolhido
            workbook.save(arquivo)

            messagebox.showinfo("Salvo com sucesso", f"Os dados foram salvos como {arquivo}")

    except Error as e:
        messagebox.showerror("Erro", f"Erro ao consultar o banco de dados: {e}")




def abrir_dados():
    nome_selecionado = cb_nomes.get()
    mes_selecionado = cb_mes.get()
    valor_mes = nome_mes_para_valor(mes_selecionado)

    try:
        con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
        consulta_sql = f'SELECT * FROM tbl_horario WHERE nomefuncionario = "{nome_selecionado}" AND MONTH(data) = "{valor_mes}"'
        cursor = con.cursor()
        cursor.execute(consulta_sql)
        dados = cursor.fetchall()

        # Criar um documento Word
        doc = docx.Document()

        # Adicionar a mensagem "Edite aqui" antes dos dados
        edit_message = "Edite aqui:\n\n"
        p1 = doc.add_paragraph()
        p1.add_run(edit_message).bold = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Adicionar os dados ao documento
        for linha in dados:
            id_funcionario, nomefuncionario, data, ini_exp_time, ini_alm_time, fim_alm_time, fim_exp_time = linha
            linha_formatada = f"Cpf: {id_funcionario}, Nome: {nomefuncionario}, Data: {data}, Início do Expediente: {ini_exp_time}, Início do Almoço: {ini_alm_time}, Término do Almoço: {fim_alm_time}, Término do Expediente: {fim_exp_time}"
            p2 = doc.add_paragraph()
            p2.add_run(linha_formatada).font.size = Pt(12)

        # Adicionar a mensagem "Edite aqui" após os dados
        p3 = doc.add_paragraph()
        p3.add_run("\n\nEdite aqui").bold = True
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Salvar o documento em um local escolhido pelo usuário
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
        if file_path:
            doc.save(file_path)

        print("Documento Word salvo com sucesso!")

    except Error as e:
        print(f"Erro ao buscar dados do funcionário: {e}")
    finally:
        if con.is_connected():
            cursor.close()
            con.close()



janela_extrato = customtkinter.CTk()
janela_extrato.title("Extrato Mensal Funcionários")
janela_extrato.geometry("270x150")
janela_extrato.resizable(width=False, height=False)
center(janela_extrato)

nomes_funcionarios = buscar_nomes_funcionarios()  # Obtém os nomes dos funcionários do banco de dados

cb_nomes = customtkinter.CTkComboBox(janela_extrato, values=nomes_funcionarios, width=250)
cb_nomes.set("Selecione o Funcionário")
cb_nomes.grid(padx=10, pady=10, row=0, column=0,columnspan=250,sticky='w')

meses_do_ano = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
]

cb_mes = customtkinter.CTkComboBox(janela_extrato, values=meses_do_ano, width=250)
cb_mes.set("Selecione o mês")
cb_mes.grid(padx=10, pady=10, row=1, column=0,columnspan=250,sticky='w')


def close():
    janela_extrato.destroy()
    admwindows()




botao_salvar_word = customtkinter.CTkButton(janela_extrato, text="Salvar Word", command=abrir_dados, width=30, height=30,
                                       border_width=2)
botao_salvar_word.grid(row=2, column=1, padx=10, pady=5,sticky="w")

botao_salvar_excel = customtkinter.CTkButton(janela_extrato, text="Salvar Excel", command=buscar_e_salvar_como_excel, width=30,
                                            height=30, border_width=2)
botao_salvar_excel.grid(row=2, column=0, padx=10, pady=5, sticky="w")

botao_voltar = customtkinter.CTkButton(janela_extrato, text="Voltar", command=close, width=30, height=30,
                                       border_width=2)
botao_voltar.grid(row=2, column=3, padx=10, pady=5, sticky="w")

janela_extrato.mainloop()