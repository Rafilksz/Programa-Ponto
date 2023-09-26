import tkinter as tk
import re
from PIL import Image, ImageTk
from tkinter import messagebox
import customtkinter
import mysql.connector
from mysql.connector import Error
from funcoes import *
import openpyxl
from openpyxl import Workbook
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from tkinter import filedialog
import datetime
from tkinter import END
from docx import Document
from docx.shared import Inches  # Importe a classe Inches


#faz janelas abrirem no meio da tela
def center(win):#faz janelas abrirem no meio da tela
    # :param win: the main window or Toplevel window to center
    # Apparently a common hack to get the window size. Temporarily hide the
    # window to avoid update_idletasks() drawing the window in the wrong
    # position.
    win.update_idletasks()  # Update "requested size" from geometry manager
    # define window dimensions width and height
    width = win.winfo_width()
    frm_width = win.winfo_rootx() - win.winfo_x()
    win_width = width + 2 * frm_width
    height = win.winfo_height()
    titlebar_height = win.winfo_rooty() - win.winfo_y()
    win_height = height + titlebar_height + frm_width
    # Get the window position from the top dynamically as well as position from left or right as follows
    x = win.winfo_screenwidth() // 2 - win_width // 2
    y = win.winfo_screenheight() // 2 - win_height // 2
    # this is the line that will center your window
    win.geometry('{}x{}+{}+{}'.format(width, height, x, y))
    # This seems to draw the window frame immediately, so only call deiconify()
    # after setting correct window position
    win.deiconify()
    # usar codigo center(NOME DA JANELA) para chamar função     #

#janela cadastro de funcionario
def novouser():
    global admin_checkbox_var
    usuario = user_id.get()
    nome = nome_id.get()
    senha = password_id.get()

    if not usuario or not nome or not senha:
        messagebox.showerror("Erro", "Preencha todos os campos obrigatórios.")
        return

    is_admin = admin_checkbox_var.get()

    # Defina o valor da coluna "adm" com base no estado da admin_checkbox
    if is_admin:
        adm_value = 1
    else:
        adm_value = 0

    try:
        con = mysql.connector.connect(host="localhost", database="db_ponto", user="root", password="D@t@W@y.1419")

        dados = f"('{usuario}', '{nome}', '{senha}', '{adm_value}')"

        add_usuario = """INSERT INTO tbl_funcionarios
        (id_funcionario, nome, senha, adm)
        VALUES """

        sql = add_usuario + dados
        print(sql)

        add_user = sql
        cursor = con.cursor()
        cursor.execute(add_user)
        con.commit()
        print(cursor.rowcount, "Registros inseridos na tabela!")
        cursor.close()

        # Limpar os campos após a inserção bem-sucedida
        user_id.delete(0, tk.END)
        nome_id.delete(0, tk.END)
        password_id.delete(0, tk.END)

    except Error as erro:
        print("falha ao inserir dados no MySQL: {}".format(erro))
    finally:
        if con.is_connected():
            con.close()
            print("conexão ao MySQL Finalizada")

def validate_integer_input(P):
    if P.isdigit():
        return True
    elif P == "":
        return True
    else:
        return False

def cadastrof():
    global user_id, nome_id, password_id, admin_checkbox_var

    customtkinter.set_appearance_mode("dark")
    customtkinter.set_default_color_theme("dark-blue")

    janela_cadastro = customtkinter.CTk()
    janela_cadastro.title("Novo Funcionário")
    janela_cadastro.geometry("300x470")
    janela_cadastro.resizable(width=False, height=False)
    center(janela_cadastro)

    Image_path = "C:\\programa ponto\\img\\New_user.png"

    def open_image2(file_path):
        image = Image.open(file_path)
        photo = ImageTk.PhotoImage(image)
        label = tk.Label(janela_cadastro, image=photo, width=150, height=150, anchor="center", background="#191919")
        label.image = photo
        label.grid(row=0, column=0, padx=10, pady=5, columnspan=300)

    open_image2(Image_path)

    def close():
        janela_cadastro.destroy()
        admwindows()

    texto = customtkinter.CTkLabel(janela_cadastro, text="Novo Cadastro", font=("Arial", 25), anchor="center")
    texto.grid(row=1, column=0, padx=10, pady=5, columnspan=300)

    txtuser = customtkinter.CTkLabel(janela_cadastro, text="CPF somente numeros")
    txtuser.grid(row=2, column=0, padx=10, pady=1, columnspan=300, sticky="w")

    vcmd = janela_cadastro.register(validate_integer_input)
    user_id = customtkinter.CTkEntry(janela_cadastro, placeholder_text="12345678910", validate="key", validatecommand=(vcmd, "%P"), width=280, height=30)
    user_id.grid(row=3, column=0, padx=10, pady=1, columnspan=300)

    txtname = customtkinter.CTkLabel(janela_cadastro, text="Nome Completo")
    txtname.grid(row=4, column=0, padx=10, pady=1, columnspan=300, sticky="w")

    nome_id = customtkinter.CTkEntry(janela_cadastro, placeholder_text="Nome Completo", width=280, height=30)
    nome_id.grid(row=5, column=0, padx=10, pady=1, columnspan=300)

    txtpass = customtkinter.CTkLabel(janela_cadastro, text="Senha")
    txtpass.grid(row=6, column=0, padx=10, pady=1, columnspan=300, sticky="w")

    password_id = customtkinter.CTkEntry(janela_cadastro, placeholder_text="Senha", width=280, height=30)
    password_id.grid(row=7, column=0, padx=1, pady=1, columnspan=300)


    admin_checkbox_var = tk.BooleanVar()
    admin_checkbox = customtkinter.CTkCheckBox(janela_cadastro, text="Administrador", variable=admin_checkbox_var,)
    admin_checkbox.grid(row=8, column=0, padx=10, pady=5, columnspan=300, sticky="w")

    botao_salvar = customtkinter.CTkButton(janela_cadastro, text="Salvar", command=novouser, width=30, height=30, border_width=2)
    botao_salvar.grid(row=9, column=0, padx=10, pady=5, columnspan=300, sticky="w")

    botao_voltar = customtkinter.CTkButton(janela_cadastro, text="Voltar", command=close, width=30, height=30, border_width=2)
    botao_voltar.grid(row=9, column=0, padx=10, pady=5, columnspan=300, sticky="e")

    janela_cadastro.mainloop()

#janela Extrato
def extratof():
    customtkinter.set_appearance_mode("dark")
    customtkinter.set_default_color_theme("dark-blue")

    # Função para converter o nome do mês em um valor numérico
    def nome_mes_para_valor(nome_mes):
        meses = {
            "Janeiro": "01",
            "Fevereiro": "02",
            "Março": "03",
            "Abril": "04",
            "Maio": "05",
            "Junho": "06",
            "Julho": "07",
            "Agosto": "08",
            "Setembro": "09",
            "Outubro": "10",
            "Novembro": "11",
            "Dezembro": "12"
        }
        return meses.get(nome_mes, "01")  # Valor padrão é "01" se o nome do mês não for encontrado

    def buscar_nomes_funcionarios():
        try:
            con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
            consulta_sql = 'SELECT nome FROM tbl_funcionarios'
            cursor = con.cursor()
            cursor.execute(consulta_sql)
            linhas = cursor.fetchall()
            nomes = [linha[0] for linha in linhas]  # Extrai os nomes da lista de tuplas

            return nomes
        except Error as e:
            print(f"Erro ao buscar nomes dos funcionários: {e}")
        finally:
            if con.is_connected():
                cursor.close()
                con.close()

    def buscar_dados_funcionario():
        nome_selecionado = cb_nomes.get()  # Obtém o nome selecionado na combobox
        mes_selecionado = cb_mes.get()  # Obtém o mês selecionado na combobox

        valor_mes = nome_mes_para_valor(mes_selecionado)  # Converte o nome do mês para valor numérico

        try:
            con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
            consulta_sql = f'SELECT * FROM tbl_horario WHERE nomefuncionario = "{nome_selecionado}" AND MONTH(data) = "{valor_mes}"'
            cursor = con.cursor()
            cursor.execute(consulta_sql)
            dados = cursor.fetchall()

            for linha in dados:
                id_funcionario, nomefuncionario, data, ini_exp_time, ini_alm_time, fim_alm_time, fim_exp_time = linha
                linha_formatada = f"Cpf: {id_funcionario},Nome {nomefuncionario}, Data: {data}, Início do Expediente: {ini_exp_time}, Início do Almoço: {ini_alm_time}, Término do Almoço: {fim_alm_time}, Término do Expediente: {fim_exp_time}"
                print(linha_formatada)

        except Error as e:
            print(f"Erro ao buscar dados do funcionário: {e}")
        finally:
            if con.is_connected():
                cursor.close()
                con.close()

    def buscar_e_salvar_como_excel():
        nome_selecionado = cb_nomes.get()

        # Obter o mês e o ano atuais
        data_atual = datetime.datetime.now()
        mes_atual = data_atual.month
        ano_atual = data_atual.year

        # Consulta SQL para buscar os dados do funcionário no mês e ano atuais
        consulta_sql = f"SELECT * FROM tbl_horario WHERE nomefuncionario = '{nome_selecionado}' AND YEAR(data) = {ano_atual} AND MONTH(data) = {mes_atual}"

        try:
            con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
            cursor = con.cursor()
            cursor.execute(consulta_sql)
            resultados = cursor.fetchall()

            # Abre uma janela de diálogo para escolher o local de salvamento
            arquivo = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Arquivos Excel", "*.xlsx")])

            if arquivo:
                # Criar um novo arquivo Excel
                workbook = Workbook()
                sheet = workbook.active

                # Definir o tamanho das colunas
                column_sizes = {'A': 15, 'B': 20, 'C': 15, 'D': 20, 'E': 15, 'F': 15, 'G': 15}
                for col, size in column_sizes.items():
                    sheet.column_dimensions[col].width = size

                # Adicionar cabeçalhos às colunas, incluindo a coluna "Soma Total"
                headers = ["CPF", "Nome Funcionario", "Data", "Inicio Expediente", "Inicio Almoço", "Fim Almoço",
                           "Fim Expediente", "Soma Total"]
                sheet.append(headers)

                # Adicionar os dados do banco de dados ao arquivo Excel e calcular a soma total
                for resultado in resultados:
                    cpf = resultado[0]
                    nome_funcionario = resultado[1]
                    data = resultado[2]
                    inicio_expediente = resultado[3]
                    inicio_almoco = resultado[4]
                    fim_almoco = resultado[5]
                    fim_expediente = resultado[6]  # Supondo que o fim de expediente está na sétima coluna

                    # Verificar se os valores não são None antes de calcular o tempo trabalhado
                    if inicio_expediente is not None and fim_expediente is not None and inicio_almoco is not None and fim_almoco is not None:
                        # Calcular o tempo trabalhado
                        tempo_trabalhado = (fim_expediente - inicio_expediente) - (fim_almoco - inicio_almoco)
                    else:
                        tempo_trabalhado = None

                    # Adicionar os dados à linha, incluindo o tempo trabalhado
                    linha = [cpf, nome_funcionario, data, inicio_expediente, inicio_almoco, fim_almoco, fim_expediente,
                             tempo_trabalhado]
                    sheet.append(linha)

                # Salvar o arquivo Excel no local escolhido
                workbook.save(arquivo)

                messagebox.showinfo("Salvo com sucesso", f"Os dados foram salvos como {arquivo}")

        except Error as e:
            messagebox.showerror("Erro", f"Erro ao consultar o banco de dados: {e}")

    def abrir_dados():
        nome_selecionado = cb_nomes.get()
        mes_selecionado = cb_mes.get()
        valor_mes = nome_mes_para_valor(mes_selecionado)

        try:
            con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
            consulta_sql = f'SELECT * FROM tbl_horario WHERE nomefuncionario = "{nome_selecionado}" AND MONTH(data) = "{valor_mes}"'
            cursor = con.cursor()
            cursor.execute(consulta_sql)
            dados = cursor.fetchall()

            # Criar um documento Word
            doc = Document()

            # Configurar as margens do documento (em polegadas)
            section = doc.sections[0]  # Obtenha a primeira seção do documento
            section.left_margin = Inches(1.0)  # Margem esquerda de 1 polegada
            section.right_margin = Inches(1.0)  # Margem direita de 1 polegada
            section.top_margin = Inches(1.0)  # Margem superior de 1 polegada
            section.bottom_margin = Inches(1.0)  # Margem inferior de 1 polegada

            # Adicionar a mensagem "Edite aqui" antes dos dados
            edit_message = (f"CENTRO EDUCACIONAL CAMPINAS EIRELI \n30.782.180/000171\nRUA JOSÉ PAULINO,1419 - CAMPINAS - SP - CEP:13013-001 - TEL:(19)3233-2815\n{cb_nomes.get()} \n")

            p1 = doc.add_paragraph()
            p1.add_run(edit_message).bold = True
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Adicionar os dados ao documento
            for linha in dados:
                id_funcionario, nomefuncionario, data, ini_exp_time, ini_alm_time, fim_alm_time, fim_exp_time = linha
                linha_formatada = f"Data: {data}, Início: {ini_exp_time}, Almoço: {ini_alm_time}, Término: {fim_alm_time}, Fim: {fim_exp_time}"
                p2 = doc.add_paragraph()
                p2.add_run(linha_formatada).font.size = Pt(12)

            # Adicionar a mensagem "Edite aqui" após os dados
            p3 = doc.add_paragraph()
            p3.add_run(f"\n\n____________________________________________\n{cb_nomes.get()}").bold = True
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Salvar o documento em um local escolhido pelo usuário
            file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Documentos Word", "*.docx")])
            if file_path:
                doc.save(file_path)

            print("Documento Word salvo com sucesso!")

        except Error as e:
            print(f"Erro ao buscar dados do funcionário: {e}")
        finally:
            if con.is_connected():
                cursor.close()
                con.close()

    janela_extrato = customtkinter.CTk()
    janela_extrato.title("Extrato Mensal Funcionários")
    janela_extrato.geometry("270x150")
    janela_extrato.resizable(width=False, height=False)
    center(janela_extrato)

    nomes_funcionarios = buscar_nomes_funcionarios()  # Obtém os nomes dos funcionários do banco de dados

    cb_nomes = customtkinter.CTkComboBox(janela_extrato, values=nomes_funcionarios, width=250)
    cb_nomes.set("Selecione o Funcionário")
    cb_nomes.grid(padx=10, pady=10, row=0, column=0, columnspan=250, sticky='w')

    meses_do_ano = [
        "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
    ]

    cb_mes = customtkinter.CTkComboBox(janela_extrato, values=meses_do_ano, width=250)
    cb_mes.set("Selecione o mês")
    cb_mes.grid(padx=10, pady=10, row=1, column=0, columnspan=250, sticky='w')

    def close():
        janela_extrato.destroy()
        admwindows()

    botao_salvar_word = customtkinter.CTkButton(janela_extrato, text="Salvar Word", command=abrir_dados, width=30,
                                                height=30,
                                                border_width=2)
    botao_salvar_word.grid(row=2, column=1, padx=10, pady=5, sticky="w")

    botao_salvar_excel = customtkinter.CTkButton(janela_extrato, text="Salvar Excel",
                                                 command=buscar_e_salvar_como_excel, width=30,
                                                 height=30, border_width=2)
    botao_salvar_excel.grid(row=2, column=0, padx=10, pady=5, sticky="w")

    botao_voltar = customtkinter.CTkButton(janela_extrato, text="Voltar", command=close, width=30, height=30,
                                           border_width=2)
    botao_voltar.grid(row=2, column=3, padx=10, pady=5, sticky="w")

    janela_extrato.mainloop()

#janela frequencia atual
def frequenciaatual():
    def salvar_como_excel():
        # Obter o nome selecionado na combobox
        funcionario_selecionado = cb_nomes.get()

        # Obter o mês e o ano atuais
        data_atual = datetime.datetime.now()
        mes_atual = data_atual.month
        ano_atual = data_atual.year

        # Consulta SQL para buscar os dados do funcionário no mês e ano atuais
        consulta_sql = f"SELECT * FROM tbl_horario WHERE nomefuncionario = '{funcionario_selecionado}' AND YEAR(data) = {ano_atual} AND MONTH(data) = {mes_atual}"

        try:
            cursor.execute(consulta_sql)
            resultados = cursor.fetchall()

            # Abre uma janela de diálogo para escolher o local de salvamento
            arquivo = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Arquivos Excel", "*.xlsx")])

            if arquivo:
                # Criar um novo arquivo Excel
                workbook = Workbook()
                sheet = workbook.active

                # Definir o tamanho das colunas
                column_sizes = {'A': 15, 'B': 20, 'C': 15, 'D': 20, 'E': 15, 'F': 15, 'G': 15}
                for col, size in column_sizes.items():
                    sheet.column_dimensions[col].width = size

                # Adicionar cabeçalhos às colunas, incluindo a coluna "Soma Total"
                headers = ["CPF", "Nome Funcionario", "Data", "Inicio Expediente", "Inicio Almoço", "Fim Almoço",
                           "Fim Expediente", "Soma Total"]
                sheet.append(headers)

                # Adicionar os dados do banco de dados ao arquivo Excel e calcular a soma total
                for resultado in resultados:
                    cpf = resultado[0]
                    nome_funcionario = resultado[1]
                    data = resultado[2]
                    inicio_expediente = resultado[3]
                    inicio_almoco = resultado[4]
                    fim_almoco = resultado[5]
                    fim_expediente = resultado[6]  # Supondo que o fim de expediente está na sétima coluna

                    # Verificar se os valores não são None antes de calcular o tempo trabalhado
                    if inicio_expediente is not None and fim_expediente is not None and inicio_almoco is not None and fim_almoco is not None:
                        # Calcular o tempo trabalhado
                        tempo_trabalhado = (fim_expediente - inicio_expediente) - (fim_almoco - inicio_almoco)
                    else:
                        tempo_trabalhado = None

                    # Adicionar os dados à linha, incluindo o tempo trabalhado
                    linha = [cpf, nome_funcionario, data, inicio_expediente, inicio_almoco, fim_almoco, fim_expediente,
                             tempo_trabalhado]
                    sheet.append(linha)

                # Salvar o arquivo Excel no local escolhido
                workbook.save(arquivo)

                messagebox.showinfo("Salvo com sucesso", f"Os dados foram salvos como {arquivo}")

        except Error as e:
            messagebox.showerror("Erro", f"Erro ao consultar o banco de dados: {e}")

    def obter_horario():
        # Obter o nome selecionado na combobox
        funcionario_selecionado = cb_nomes.get()

        # Obter o mês e o ano atuais
        data_atual = datetime.datetime.now()
        mes_atual = data_atual.month
        ano_atual = data_atual.year

        # Consulta SQL para buscar os dados do funcionário no mês e ano atuais
        consulta_sql = f"SELECT * FROM tbl_horario WHERE nomefuncionario = '{funcionario_selecionado}' AND YEAR(data) = {ano_atual} AND MONTH(data) = {mes_atual}"

        try:
            cursor.execute(consulta_sql)
            resultados = cursor.fetchall()

            # Limpar o conteúdo atual do widget de Text
            resultados_text.delete(1.0, "end")

            # Processar os resultados e inseri-los no widget de Text
            for resultado in resultados:
                # Converter a linha de resultado em uma string formatada
                resultado_str = ", ".join(map(str, resultado)) + "\n"
                resultados_text.insert("end", resultado_str)

        except Error as e:
            messagebox.showerror("Erro", f"Erro ao consultar o banco de dados: {e}")

    customtkinter.set_appearance_mode("dark")
    customtkinter.set_default_color_theme("dark-blue")

    # conecto banco de dados
    con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')

    consulta_sql = 'select nome from tbl_funcionarios'  # seleciono a column que vou precisar
    cursor = con.cursor()
    cursor.execute(consulta_sql)
    linhas = cursor.fetchall()

    print(f"numero total de registros retornados:  {cursor.rowcount}")

    print('\nAs Linhas são:')

    string = '-'.join(
        map(str, linhas))  # AKI EU TO PEGANDO O ARRAY QUE VEM DO BANCO E TRANFORMANDO EM STRING SEPARANDO POR -
    string = re.sub('[^a-zA-Z0-9\- \\\]', '',
                    string)  # LIMPO TODOS CARACTERES DEIXANDO APENAS LETRAS E NUMEROS E O - TAMBEM
    arr = string.split("-")  # CRIO UM NOVO ARRAY SEPARANDO CADA INDEX PELO -

    def close():
        janela_f_atual.destroy()
        admwindows()

    janela_f_atual = customtkinter.CTk()
    lado, cima = (janela_f_atual.winfo_screenwidth()), (janela_f_atual.winfo_screenheight())
    janela_f_atual.geometry('%dx%d+0+0' % (lado, cima))
    janela_f_atual.title("Frequencia Atual")
    janela_f_atual.resizable(width=False, height=False)
    listanome = arr
    texto = "meu texto aqui"

    # ComboBox
    cb_nomes = customtkinter.CTkComboBox(janela_f_atual, values=listanome, width=250)
    cb_nomes.set("Selecione o Funcionário")
    cb_nomes.grid(row=0, column=0, padx=20, pady=5, columnspan=290, sticky="w")  # Alinhado à esquerda

    # Botão
    botao_buscar = customtkinter.CTkButton(
        janela_f_atual, text="Abrir", command=obter_horario,
        width=30, height=30, border_width=2, )
    botao_buscar.grid(row=0, column=0, padx=280, pady=5, columnspan=290, sticky="w")

    # Obtém a quantidade de colunas disponíveis na grade da janela
    num_colunas = janela_f_atual.grid_size()[0]

    # Crie um widget de Text para exibir os resultados
    resultados_text = customtkinter.CTkTextbox(janela_f_atual, width=1330, height=630)
    resultados_text.grid(row=1, column=0, padx=20, pady=5, columnspan=num_colunas, sticky="w")

    botao_voltar = customtkinter.CTkButton(
        janela_f_atual, text="Voltar", command=close,
        width=30, height=30, border_width=2, )
    botao_voltar.grid(row=3, column=0, padx=20, pady=5, sticky="e")

    botao_salvar_excel = customtkinter.CTkButton(
        janela_f_atual, text="Salvar como Excel", command=salvar_como_excel,
        width=30, height=30, border_width=2)
    botao_salvar_excel.grid(row=3, column=1, padx=5, pady=10)

    janela_f_atual.mainloop()

#janela mudanca
def mudancaf():
     customtkinter.set_appearance_mode("dark")
     customtkinter.set_default_color_theme("dark-blue")
     
     def buscar_nomes_funcionarios():
         try:
             con = mysql.connector.connect(host='localhost', database='db_ponto', user='root', password='D@t@W@y.1419')
             consulta_sql = 'SELECT nome FROM tbl_funcionarios'
             cursor = con.cursor()
             cursor.execute(consulta_sql)
             linhas = cursor.fetchall()
             nomes = [linha[0] for linha in linhas]  # Extrai os nomes da lista de tuplas
     
             return nomes
         except Error as e:
             print(f"Erro ao buscar nomes dos funcionários: {e}")
         finally:
             if con.is_connected():
                 cursor.close()
                 con.close()
     
     def close():
         janela_mudanca.destroy()
         admwindows()
         con.close()
         cursor.close()
         print('conexao mysql fechada')
     
     def salvar():
         # Obter os valores selecionados
         nome_funcionario = nomes_fun.get()
         data_selecionada = data_new.get()
         hora_selecionada = hora_new.get()
         horario_selecionado = cb_horario.get()
     
         # Converter o formato da data de DD-MM-YYYY para YYYY-MM-DD
         data_formatada = datetime.datetime.strptime(data_selecionada, '%d-%m-%Y').strftime('%Y-%m-%d')
     
         # Resto do código permanece o mesmo...
     
         # Mapear os valores de cb_horario para as colunas correspondentes
         colunas_horario = {
             'Inicio Expediente': 'ini_exp_time',
             'Inicio Almoço': 'ini_alm_time',
             'Fim Almoço': 'fim_alm_time',
             'Fim Expediente': 'fim_exp_time'
         }
     
         coluna_horario = colunas_horario.get(horario_selecionado, '')
     
         if nome_funcionario and data_formatada and coluna_horario:
             try:
                 con = mysql.connector.connect(host='localhost', database='db_ponto', user='root',
                                               password='D@t@W@y.1419')
                 cursor = con.cursor()
     
                 # Verificar se já existe um registro com os mesmos valores de nome_funcionario e data
                 select_sql = "SELECT * FROM tbl_horario WHERE nomefuncionario = %s AND data = %s"
                 cursor.execute(select_sql, (nome_funcionario, data_formatada))
                 resultado = cursor.fetchone()
     
                 if resultado:
                     # Se o registro já existe, atualize a coluna correspondente
                     update_sql = f"UPDATE tbl_horario SET {coluna_horario} = %s WHERE nomefuncionario = %s AND data = %s"
                     cursor.execute(update_sql, (hora_selecionada, nome_funcionario, data_formatada))
                 else:
                     # Se o registro não existe, insira um novo registro com os valores adequados
                     insert_sql = "INSERT INTO tbl_horario (nomefuncionario, data, " + coluna_horario + ") VALUES (%s, %s, %s)"
                     cursor.execute(insert_sql, (nome_funcionario, data_formatada, hora_selecionada))
     
                 con.commit()
                 messagebox.showinfo("Sucesso", "Registro atualizado com sucesso!")
     
             except mysql.connector.Error as e:
                 messagebox.showerror("Erro", f"Erro ao atualizar registro: {e}")
             finally:
                 if con.is_connected():
                     cursor.close()
                     con.close()
         else:
             messagebox.showerror("Erro", "Por favor, selecione um funcionário, data, hora e horário válidos.")
     
     # Resto do seu código...
     
     def formatar_data(event):
         data = data_new.get()
         if len(data) == 2:
             data_new.delete(0, END)
             data_new.insert(0, data + '-')
         elif len(data) == 5:
             data_new.delete(0, END)
             data_new.insert(0, data + '-')
         elif len(data) == 10:
             hora_new.focus_set()  # Set focus to hora_new field
     
     # Função para formatar a entrada de hora no estilo HH:MM
     def formatar_hora(event):
         hora = hora_new.get()
     
         # Remova qualquer caractere que não seja um dígito
         hora = ''.join(filter(str.isdigit, hora))
     
         if len(hora) == 2:
             hora_new.delete(0, END)
             hora_new.insert(0, hora + ':')
         elif len(hora) == 4:
             hora_new.delete(0, END)
             hora_new.insert(0, hora[:2] + ':' + hora[2:] + ':00')
             cb_horario.focus_set()  # Definir o foco para cb_horario
         elif len(hora) == 8:
             hora_new.delete(0, END)
             hora_new.insert(0, hora[:5] + ':00')
             cb_horario.focus_set()  # Definir o foco para cb_horario
     
     janela_mudanca = customtkinter.CTk()
     janela_mudanca.title("Controle de Ponto Dataway")
     janela_mudanca.geometry("260x160")
     janela_mudanca.resizable(width=False, height=False)
     center(janela_mudanca)
     
     listhr = ['Inicio Expediente', 'Inicio Almoço', 'Fim Almoço', 'Fim Expediente']
     
     nomes_funcionarios = buscar_nomes_funcionarios()  # Obtém os nomes dos funcionários do banco de dados
     nomes_fun = customtkinter.CTkComboBox(master=janela_mudanca, values=nomes_funcionarios, width=250)
     nomes_fun.set("Selecione o Funcionário")
     nomes_fun.grid(row=0, column=0, padx=5, pady=5, )
     
     data_new = customtkinter.CTkEntry(janela_mudanca, placeholder_text="DD/MM/AAAA", width=100, height=30, )
     data_new.grid(row=1, column=0, padx=5, pady=5, sticky="w")
     data_new.bind('<KeyRelease>', formatar_data)
     
     hora_new = customtkinter.CTkEntry(janela_mudanca, placeholder_text="HH:MM", width=60, height=30, )
     hora_new.grid(row=1, column=0, padx=5, pady=5, sticky="e")
     
     data_new.bind('<KeyRelease>', formatar_data)
     hora_new.bind('<KeyRelease>', formatar_hora)
     
     cb_horario = customtkinter.CTkComboBox(janela_mudanca, values=listhr, width=250)
     (cb_horario.set("Horario esquecido"))
     cb_horario.grid(row=2, column=0, padx=5, pady=5, )
     
     botao_voltar = customtkinter.CTkButton(
         janela_mudanca, text="Voltar", command=close,
         width=30, height=30, border_width=2, )
     botao_voltar.grid(row=3, column=0, padx=5, pady=5, sticky="e")
     
     botao_salvar = customtkinter.CTkButton(
         janela_mudanca, text="Salvar", command=salvar,
         width=30, height=30, border_width=2, )
     botao_salvar.grid(row=3, column=0, padx=5, pady=5, )
     
     janela_mudanca.mainloop()


#menu principal ADM
def admwindows():

    customtkinter.set_appearance_mode("dark")
    customtkinter.set_default_color_theme("dark-blue")
    janela_adm = customtkinter.CTk()
    janela_adm.geometry("250x250")
    janela_adm.title("Menu Administrador")
    janela_adm.resizable(width=False, height=False)
    center(janela_adm)

    def novof():
        janela_adm.destroy()
        cadastrof()

    def alterar():
        janela_adm.destroy()
        mudancaf()

    def consultar():
        janela_adm.destroy()
        frequenciaatual()

    def extrato():
        janela_adm.destroy()
        extratof()

    # Defina o caminho para a imagem
    newuserp = "C:\\programa ponto\\img\\New_userp.png"
    alterahr = "C:\\programa ponto\\img\\altera.png"
    consultap = "C:\\programa ponto\\img\\consultar.png"
    extratop = "C:\\programa ponto\\img\\extrato.png"

    def open_image3(newuserp, width, height):
        image3 = Image.open(newuserp)
        image3 = image3.resize((width, height))
        photonu = ImageTk.PhotoImage(image3)

        return photonu

    def open_image4(alterahr, width, height):
        image4 = Image.open(alterahr)
        image4 = image4.resize((width, height))
        altera = ImageTk.PhotoImage(image4)

        return altera

    def open_image5(consultap, width, height):
        image5 = Image.open(consultap)
        image5 = image5.resize((width, height))
        consulta = ImageTk.PhotoImage(image5)

        return consulta

    def open_image6(extratop, width, height):
        image6 = Image.open(extratop)
        image6 = image6.resize((width, height))
        extratoi = ImageTk.PhotoImage(image6)

        return extratoi


    # Defina o seu botão personalizado com a imagem
    photonu = open_image3(newuserp, width=89, height=89)
    botao_cadastro = customtkinter.CTkButton(janela_adm, text="",command=novof, image=photonu, width=90, height=90)
    botao_cadastro.grid(row=0, column=0, padx=10, pady=10, )

    altera = open_image4(alterahr, width=89, height=89)
    altera_hora = customtkinter.CTkButton(janela_adm, text="",command=alterar, image=altera, width=90, height=90)
    altera_hora.grid(row=0, column=1, padx=10, pady=10, )

    consulta = open_image5(consultap, width=89, height=89)
    botao_consulta_freq = customtkinter.CTkButton(janela_adm, text="",command=consultar, image=consulta, width=90, height=90)
    botao_consulta_freq.grid(row=1, column=0, padx=10, pady=10, )

    extratoi = open_image6(extratop, width=89, height=89)
    botao_extrato_hr = customtkinter.CTkButton(janela_adm, text="",command=extrato, image=extratoi, width=90, height=90)
    botao_extrato_hr.grid(row=1, column=1, padx=10, pady=10, )

    janela_adm.mainloop()

#janela bater ponto


